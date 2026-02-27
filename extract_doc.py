from docx import Document
import os

base_path = r"c:\GitWorkspace\Mage Promo\event-logger-ms\docs"

# Извлечение из ТЗ
print("=" * 60)
print("Event Logger - ТЗ.docx")
print("=" * 60)

doc_tz = Document(os.path.join(base_path, "Event Logger - ТЗ.docx"))

full_text = []
for para in doc_tz.paragraphs:
    if para.text.strip():
        full_text.append(para.text)

for table in doc_tz.tables:
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            if cell.text.strip():
                row_text.append(cell.text.strip())
        if row_text:
            full_text.append(" | ".join(row_text))

print("\n".join(full_text))

print("\n" + "=" * 60)
print("Event Logger - события системы.docx")
print("=" * 60)

doc_events = Document(os.path.join(base_path, "Event Logger - события системы.docx"))

full_text = []
for para in doc_events.paragraphs:
    if para.text.strip():
        full_text.append(para.text)

for table in doc_events.tables:
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            if cell.text.strip():
                row_text.append(cell.text.strip())
        if row_text:
            full_text.append(" | ".join(row_text))

print("\n".join(full_text))
